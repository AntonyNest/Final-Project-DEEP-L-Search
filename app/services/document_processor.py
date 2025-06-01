# Document parsing logic

"""
Document Processor Service - Екстракція та підготовка тексту з документів.

Цей сервіс реалізує шаблон "Extract-Transform-Load" (ETL) для документів.
Він відповідає за: 
1) екстракцію тексту з різних форматів
2) очищення та нормалізацію тексту  
3) розбиття на семантично значущі чанки
4) екстракцію метаданих
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import logging

# Document parsing libraries
from docx import Document
import PyPDF2

# Our configuration
from app.config import settings

logger = logging.getLogger(__name__)


class DocumentChunk:
    """
    Модель для зберігання обробленого чанку документа.
    
    Ця модель використовує composition pattern - вона містить не лише текст,
    але й контекстну інформацію, яка критично важлива для якісного пошуку.
    """
    
    def __init__(
        self, 
        text: str, 
        chunk_id: str,
        source_file: str,
        page_number: Optional[int] = None,
        chunk_index: int = 0,
        metadata: Optional[Dict[str, Any]] = None
    ):
        self.text = text
        self.chunk_id = chunk_id
        self.source_file = source_file
        self.page_number = page_number
        self.chunk_index = chunk_index
        self.metadata = metadata or {}
        
        # Автоматично генеруємо додаткові метадані
        self.word_count = len(text.split())
        self.char_count = len(text)
        self.created_at = datetime.now().isoformat()
    
    def to_dict(self) -> Dict[str, Any]:
        """Конвертує чанк в словник для зберігання в векторній БД."""
        return {
            "text": self.text,
            "chunk_id": self.chunk_id,
            "source_file": self.source_file,
            "page_number": self.page_number,
            "chunk_index": self.chunk_index,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "created_at": self.created_at,
            **self.metadata
        }


class DocumentProcessor:
    """
    Головний клас для обробки документів.
    
    Архітектурний підхід: Single Responsibility Principle - кожний метод
    відповідає за одну конкретну задачу в процесі обробки документів.
    """
    
    def __init__(self):
        """Ініціалізація процесора з налаштуваннями з конфігурації."""
        self.chunk_size = settings.max_chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.documents_path = settings.documents_path_obj
        
        # Підтримувані формати файлів
        self.supported_extensions = {'.docx', '.doc', '.pdf', '.txt'}
        
        logger.info(
            f"DocumentProcessor initialized with chunk_size={self.chunk_size}, "
            f"overlap={self.chunk_overlap}, path={self.documents_path}"
        )
    
    def discover_documents(self) -> List[Path]:
        """
        Знаходить всі підтримувані документи в директорії.
        
        Використовує рекурсивний пошук для обходу всіх піддиректорій.
        Це важливо для великих архівів документів із складною структурою.
        """
        documents = []
        
        try:
            # Рекурсивний обхід всіх файлів
            for file_path in self.documents_path.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                    documents.append(file_path)
                    
            logger.info(f"Discovered {len(documents)} documents")
            return documents
            
        except Exception as e:
            logger.error(f"Error discovering documents: {str(e)}")
            return []
    
    def extract_text_from_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Екстракція тексту з .docx файлу з використанням python-docx.
        
        Returns:
            tuple: (extracted_text, metadata_dict)
        """
        try:
            doc = Document(file_path)
            
            # Екстракція тексту з параграфів
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Пропускаємо порожні параграфи
                    paragraphs.append(text)
            
            # Екстракція тексту з таблиць (важливо для структурованих документів)
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            # Об'єднуємо весь текст
            full_text = "\n".join(paragraphs)
            if table_text:
                full_text += "\n\nТаблиці:\n" + "\n".join(table_text)
            
            # Екстракція метаданих з властивостей документа
            metadata = {
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            }
            
            # Додаємо метадані з властивостей документа, якщо вони доступні
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                metadata["created_date"] = doc.core_properties.created.isoformat()
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return "", {}
    
    def extract_text_from_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Робастна екстракція тексту з PDF файлу з fallback стратегією.
        
        1. Спочатку PyPDF2 (швидко, для простих PDF)
        2. Fallback до pdfplumber (для складних макетів)
        3. Додаткові метадані про метод екстракції
        """
        metadata = {"file_type": "pdf"}
        
        # Метод 1: PyPDF2 - найшвидший для простих PDF
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text_parts = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Сторінка {page_num + 1}]\n{page_text}")
                
                full_text = "\n\n".join(text_parts)
                
                # Базові метадані
                metadata.update({
                    "page_count": len(pdf_reader.pages),
                    "extraction_method": "PyPDF2"
                })
                
                # Екстракція метаданих PDF
                if pdf_reader.metadata:
                    if pdf_reader.metadata.get('/Title'):
                        metadata["title"] = str(pdf_reader.metadata['/Title'])
                    if pdf_reader.metadata.get('/Author'):
                        metadata["author"] = str(pdf_reader.metadata['/Author'])
                    if pdf_reader.metadata.get('/CreationDate'):
                        metadata["created_date"] = str(pdf_reader.metadata['/CreationDate'])
                
                # Якщо витягнули достатньо тексту, повертаємо результат
                if len(full_text.strip()) > 50:  # Мінімальна кількість символів
                    logger.debug(f"PDF extracted with PyPDF2: {len(full_text)} chars")
                    return full_text, metadata
                else:
                    logger.warning(f"PyPDF2 extracted insufficient text ({len(full_text)} chars), trying fallback")
                    
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_path}: {str(e)}, trying fallback")
        
        # Метод 2: Fallback до pdfplumber для складних PDF
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                
                for page_num, page in enumerate(pdf.pages):
                    # Витягуємо звичайний текст
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Сторінка {page_num + 1}]\n{page_text}")
                    
                    # Витягуємо таблиці окремо (важливо для структурованих документів)
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            table_text = f"[Таблиця {table_num + 1} на сторінці {page_num + 1}]\n"
                            for row in table:
                                if row and any(cell for cell in row if cell):
                                    row_text = " | ".join([str(cell or "") for cell in row])
                                    table_text += row_text + "\n"
                            text_parts.append(table_text)
                
                full_text = "\n\n".join(text_parts)
                metadata.update({
                    "page_count": len(pdf.pages),
                    "extraction_method": "pdfplumber",
                    "has_tables": any(page.extract_tables() for page in pdf.pages)
                })
                
                if len(full_text.strip()) > 10:
                    logger.info(f"PDF extracted with pdfplumber: {len(full_text)} chars")
                    return full_text, metadata
                
        except ImportError:
            logger.warning("pdfplumber not available, skipping fallback")
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {str(e)}")
        
        # Якщо всі методи не спрацювали
        logger.error(f"All PDF extraction methods failed for {file_path}")
        metadata["extraction_method"] = "failed"
        return "", metadata
    
    def extract_text_from_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Екстракція тексту з простого текстового файлу."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "file_type": "txt",
                "line_count": len(text.split('\n')),
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Спробуємо з різними кодуваннями
            for encoding in ['cp1251', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    metadata = {"file_type": "txt", "encoding": encoding}
                    return text, metadata
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode text file {file_path}")
            return "", {}
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return "", {}
    
    def clean_text(self, text: str) -> str:
        """
        Очищення та нормалізація тексту.
        
        Цей метод реалізує preprocessing pipeline для покращення якості
        векторизації та пошуку. Кожна операція має конкретне обґрунтування.
        """
        # Заміна множинних пробілів та переносів рядків
        text = re.sub(r'\s+', ' ', text)
        
        # Видалення зайвих символів, але зберігання пунктуації (важливо для семантики)
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\']+', ' ', text, flags=re.UNICODE)
        
        # Видалення множинних розділових знаків
        text = re.sub(r'[\.]{2,}', '.', text)
        text = re.sub(r'[\!\?]{2,}', '!', text)
        
        # Прибираємо зайві пробіли
        text = text.strip()
        
        return text
    
    def split_into_chunks(self, text: str, metadata: Dict[str, Any]) -> List[str]:
        """
        Розбиття тексту на чанки з overlap для збереження контексту.
        
        Алгоритм:
        1. Намагаємося розбивати по реченнях (семантично значущі границі)
        2. Якщо речення довше за максимальний розмір, розбиваємо по словах
        3. Додаємо overlap між чанками для збереження контексту
        """
        if not text or len(text) < self.chunk_size:
            return [text] if text else []
        
        # Спочатку розбиваємо по реченнях
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # Якщо додавання речення не перевищує розмір чанку
            if len(current_chunk + " " + sentence) <= self.chunk_size:
                current_chunk += " " + sentence if current_chunk else sentence
            else:
                # Зберігаємо поточний чанк, якщо він не порожній
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Починаємо новий чанк
                if len(sentence) <= self.chunk_size:
                    current_chunk = sentence
                else:
                    # Якщо речення довше за максимальний розмір, розбиваємо по словах
                    words = sentence.split()
                    word_chunk = ""
                    
                    for word in words:
                        if len(word_chunk + " " + word) <= self.chunk_size:
                            word_chunk += " " + word if word_chunk else word
                        else:
                            if word_chunk:
                                chunks.append(word_chunk.strip())
                            word_chunk = word
                    
                    current_chunk = word_chunk
        
        # Додаємо останній чанк
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Додаємо overlap між чанками для збереження контексту
        overlapped_chunks = []
        for i, chunk in enumerate(chunks):
            if i > 0 and self.chunk_overlap > 0:
                # Берємо останні слова з попереднього чанку
                prev_words = chunks[i-1].split()[-self.chunk_overlap//10:]  # Приблизно
                overlap_text = " ".join(prev_words)
                chunk = overlap_text + " " + chunk
            
            overlapped_chunks.append(chunk)
        
        return overlapped_chunks
    
    def process_document(self, file_path: Path) -> List[DocumentChunk]:
        """
        Головний метод обробки документа.
        
        Повертає список об'єктів DocumentChunk, готових для векторизації.
        """
        logger.info(f"Processing document: {file_path}")
        
        # Екстракція тексту на основі типу файлу
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.docx':
            text, doc_metadata = self.extract_text_from_docx(file_path)
        elif file_extension == '.pdf':
            text, doc_metadata = self.extract_text_from_pdf(file_path)
        elif file_extension == '.txt':
            text, doc_metadata = self.extract_text_from_txt(file_path)
        elif file_extension == '.doc':
            # Для .doc файлів потрібен додатковий інструментарій
            logger.warning(f"Legacy .doc format not fully supported: {file_path}")
            return []
        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            return []
        
        if not text:
            logger.warning(f"No text extracted from: {file_path}")
            return []
        
        # Очищення тексту
        cleaned_text = self.clean_text(text)
        
        # Розбиття на чанки
        text_chunks = self.split_into_chunks(cleaned_text, doc_metadata)
        
        # Створення об'єктів DocumentChunk
        document_chunks = []
        for i, chunk_text in enumerate(text_chunks):
            # Генеруємо унікальний ID для чанку
            chunk_id = f"{file_path.stem}_{i:04d}"
            
            # Додаємо метадані файлу
            chunk_metadata = {
                **doc_metadata,
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "file_modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
            }
            
            chunk = DocumentChunk(
                text=chunk_text,
                chunk_id=chunk_id,
                source_file=str(file_path),
                chunk_index=i,
                metadata=chunk_metadata
            )
            
            document_chunks.append(chunk)
        
        logger.info(f"Created {len(document_chunks)} chunks from {file_path}")
        return document_chunks
    
    def process_all_documents(self) -> List[DocumentChunk]:
        """
        Обробка всіх документів в директорії.
        
        Цей метод координує обробку всіх знайдених документів та 
        повертає загальний список чанків для індексації.
        """
        all_chunks = []
        documents = self.discover_documents()
        
        logger.info(f"Starting processing of {len(documents)} documents")
        
        for doc_path in documents:
            try:
                chunks = self.process_document(doc_path)
                all_chunks.extend(chunks)
            except Exception as e:
                logger.error(f"Failed to process {doc_path}: {str(e)}")
                continue
        
        logger.info(f"Processed {len(documents)} documents into {len(all_chunks)} chunks")
        return all_chunks